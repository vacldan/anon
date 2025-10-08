# -*- coding: utf-8 -*-
"""
Simple Czech Document Anonymizer
================================
A focused anonymization system for Czech documents that detects and replaces
sensitive personal data with anonymized markers.

Usage:
    python anonymizer_simple.py input.docx [--output output.docx]
"""

import sys
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Set, Tuple
from dataclasses import dataclass
import argparse
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ========== Configuration ==========

@dataclass
class AnonymizationResult:
    """Result of anonymization process"""
    original_text: str
    anonymized_text: str
    replacements: Dict[str, List[str]]
    statistics: Dict[str, int]
    processing_time: float

# ========== Utility Functions ==========

def normalize_text(text: str) -> str:
    """Normalize text for consistent processing"""
    return text.strip()

def setup_logging() -> logging.Logger:
    """Setup logging configuration"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('anonymizer.log', encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger(__name__)

# ========== Czech Name Detection ==========

class CzechNameDetector:
    """Czech name detection with common names database"""
    
    def __init__(self):
        self.male_names = {
            'jan', 'petr', 'pavel', 'tomáš', 'martin', 'jaroslav', 'milan', 'františek',
            'josef', 'antonín', 'zdeněk', 'vladimír', 'stanislav', 'luděk', 'karel',
            'michal', 'david', 'lukáš', 'ondřej', 'jakub', 'matěj', 'adam', 'daniel',
            'filip', 'mikuláš', 'vít', 'matyáš', 'kryštof', 'sebastian', 'benjamin',
            'ondra', 'honza', 'pepa', 'míra', 'jirka', 'kuba', 'tonda', 'václav'
        }
        
        self.female_names = {
            'marie', 'jana', 'eva', 'hana', 'anna', 'věra', 'alena', 'lenka',
            'kateřina', 'lucie', 'petra', 'zuzana', 'iveta', 'monika', 'veronika',
            'tereza', 'barbora', 'adéla', 'karolína', 'kristýna', 'nikola', 'natálie',
            'eliska', 'sophie', 'emma', 'olivia', 'amélie', 'aneta', 'klára', 'julie'
        }
        
        self.surnames = {
            'novák', 'svoboda', 'novotný', 'dvořák', 'černý', 'procházka',
            'kučera', 'veselý', 'horák', 'němec', 'pokorný', 'pospíšil',
            'havel', 'bláha', 'krejčí', 'stárek', 'kříž', 'beneš', 'fiala',
            'moravec', 'barták', 'urban', 'polák', 'doležal', 'šimánek',
            'nováková', 'svobodová', 'novotná', 'dvořáková', 'černá', 'procházková'
        }
        
        # Common surname suffixes
        self.surname_suffixes = {'ová', 'ek', 'ík', 'ák', 'ček', 'čík', 'ko', 'ka', 'ja', 'ský', 'cký'}
    
    def is_likely_first_name(self, word: str) -> bool:
        """Check if word is likely a first name"""
        normalized = word.lower().strip()
        return normalized in self.male_names or normalized in self.female_names
    
    def is_likely_surname(self, word: str) -> bool:
        """Check if word is likely a surname"""
        normalized = word.lower().strip()
        return (normalized in self.surnames or 
                any(normalized.endswith(suffix) for suffix in self.surname_suffixes))

# ========== Pattern Detection ==========

class PatternDetector:
    """Pattern detection for sensitive data"""
    
    def __init__(self):
        self.patterns = [
            # Birth dates (more specific)
            (r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b', 'DATE'),
            # Czech birth number (RČ) - more specific
            (r'\b\d{2}[0156]\d{3,4}/\d{4}\b', 'BIRTH_ID'),
            # ID card number (with context)
            (r'\b\d{9}\b', 'ID_CARD'),
            # Bank account (more specific)
            (r'\b\d{1,6}-\d{1,10}/\d{4}\b', 'BANK'),
            # IBAN
            (r'\bCZ\d{2}(?:\s?\d){20}\b', 'BANK'),
            # Phone numbers
            (r'(?:\+?420[\s\-]?)?(?<!\d)(?:\d{3}[\s\-]?){2}\d{3}(?!\d)', 'PHONE'),
            # Email addresses
            (r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', 'EMAIL'),
            # Addresses (more specific)
            (r'\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]{2,30}\s+\d{1,4}(?:/\d{1,4})?,\s*\d{3}\s?\d{2}\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]{2,20}\b', 'ADDRESS'),
        ]
    
    def detect_patterns(self, text: str) -> List[Tuple[str, str, int, int]]:
        """Detect all patterns in text and return (category, match, start, end)"""
        results = []
        
        for pattern, category in self.patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                # Skip if it looks like a legal reference
                context = text[max(0, match.start()-20):match.end()+20].lower()
                if any(x in context for x in ['§', 'zákon', 'oz', 'vyhláška']):
                    continue
                
                # Special validation for ID card numbers
                if category == 'ID_CARD':
                    if not any(x in context for x in ['op', 'občansk', 'průkaz']):
                        continue
                
                results.append((category, match.group(), match.start(), match.end()))
        
        return results

# ========== Main Anonymizer Class ==========

class SimpleAnonymizer:
    """Simple anonymizer with focused detection"""
    
    def __init__(self):
        self.name_detector = CzechNameDetector()
        self.pattern_detector = PatternDetector()
        self.logger = logging.getLogger(__name__)
        
        # Mapping and tracking
        self.replacements: Dict[str, List[str]] = {}
        self.counters: Dict[str, int] = {}
        self.person_mappings: Dict[Tuple[str, str], str] = {}
    
    def _new_tag(self, category: str) -> str:
        """Generate new anonymization tag"""
        self.counters[category] = self.counters.get(category, 0) + 1
        return f"[[{category}_{self.counters[category]}]]"
    
    def _add_replacement(self, tag: str, original: str) -> None:
        """Add replacement to mapping"""
        if tag not in self.replacements:
            self.replacements[tag] = []
        if original not in self.replacements[tag]:
            self.replacements[tag].append(original)
    
    def _replace_text(self, text: str, start: int, end: int, tag: str, original: str) -> str:
        """Replace text span with anonymization tag"""
        if start < 0 or end > len(text) or start >= end:
            return text
        
        # Check if already anonymized
        if text[start:end].startswith("[[") and text[start:end].endswith("]]"):
            return text
        
        self._add_replacement(tag, original)
        return text[:start] + tag + text[end:]
    
    def anonymize_names(self, text: str) -> str:
        """Anonymize names using simple detection"""
        # Find potential names (capitalized words)
        name_pattern = re.compile(r'\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+\b')
        
        def replace_name(match):
            word = match.group()
            cleaned = re.sub(r'[^\w]', '', word)
            
            # Check if it's likely a name
            if (self.name_detector.is_likely_first_name(cleaned) or 
                self.name_detector.is_likely_surname(cleaned)):
                
                # Look for name pairs
                start, end = match.span()
                next_text = text[end:end+100]  # Look ahead
                next_words = re.findall(r'\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+\b', next_text)
                
                for next_word in next_words[:2]:  # Check next 2 words
                    next_cleaned = re.sub(r'[^\w]', '', next_word)
                    if self.name_detector.is_likely_surname(next_cleaned):
                        # Found name pair
                        key = (cleaned.lower(), next_cleaned.lower())
                        if key not in self.person_mappings:
                            tag = self._new_tag("PERSON")
                            self.person_mappings[key] = tag
                        tag = self.person_mappings[key]
                        self._add_replacement(tag, f"{word} {next_word}")
                        return tag
                
                # Single name
                key = (cleaned.lower(), "")
                if key not in self.person_mappings:
                    tag = self._new_tag("PERSON")
                    self.person_mappings[key] = tag
                tag = self.person_mappings[key]
                self._add_replacement(tag, word)
                return tag
            
            return word
        
        return name_pattern.sub(replace_name, text)
    
    def anonymize_patterns(self, text: str) -> str:
        """Anonymize detected patterns"""
        patterns = self.pattern_detector.detect_patterns(text)
        
        # Sort by position (reverse order to avoid offset issues)
        patterns.sort(key=lambda x: x[2], reverse=True)
        
        # Track processed ranges to avoid conflicts
        processed_ranges = []
        
        for category, match, start, end in patterns:
            # Check if this range overlaps with already processed ranges
            overlaps = any(not (end <= existing_start or start >= existing_end) 
                          for existing_start, existing_end in processed_ranges)
            
            if not overlaps:
                tag = self._new_tag(category)
                text = self._replace_text(text, start, end, tag, match)
                processed_ranges.append((start, end))
        
        return text
    
    def anonymize_text(self, text: str) -> str:
        """Main anonymization method"""
        if not text.strip():
            return text
        
        # Anonymize names first
        text = self.anonymize_names(text)
        
        # Anonymize patterns
        text = self.anonymize_patterns(text)
        
        return text
    
    def get_statistics(self) -> Dict[str, int]:
        """Get anonymization statistics"""
        return dict(self.counters)

# ========== Document Processing ==========

class DocumentProcessor:
    """Process various document formats"""
    
    def __init__(self, anonymizer: SimpleAnonymizer):
        self.anonymizer = anonymizer
        self.logger = logging.getLogger(__name__)
    
    def process_docx(self, input_path: Path, output_path: Path) -> AnonymizationResult:
        """Process DOCX document"""
        start_time = datetime.now()
        
        try:
            doc = Document(str(input_path))
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original = paragraph.text
                    anonymized = self.anonymizer.anonymize_text(original)
                    paragraph.text = anonymized
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original = paragraph.text
                                anonymized = self.anonymizer.anonymize_text(original)
                                paragraph.text = anonymized
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return AnonymizationResult(
                original_text="",  # Not stored for large documents
                anonymized_text="",  # Not stored for large documents
                replacements=self.anonymizer.replacements,
                statistics=self.anonymizer.get_statistics(),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
    
    def process_text(self, input_path: Path, output_path: Path) -> AnonymizationResult:
        """Process plain text document"""
        start_time = datetime.now()
        
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                original_text = f.read()
            
            anonymized_text = self.anonymizer.anonymize_text(original_text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(anonymized_text)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return AnonymizationResult(
                original_text=original_text,
                anonymized_text=anonymized_text,
                replacements=self.anonymizer.replacements,
                statistics=self.anonymizer.get_statistics(),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Error processing text file: {e}")
            raise

# ========== Output and Mapping ==========

def save_mappings(base_path: Path, replacements: Dict[str, List[str]], statistics: Dict[str, int]):
    """Save anonymization mappings in multiple formats"""
    
    # JSON format
    json_path = base_path.with_suffix('.json')
    mapping_data = {
        'metadata': {
            'created_at': datetime.now().isoformat(),
            'statistics': statistics,
            'total_replacements': sum(len(vals) for vals in replacements.values())
        },
        'replacements': replacements
    }
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(mapping_data, f, ensure_ascii=False, indent=2)
    
    # Text format
    txt_path = base_path.with_suffix('.txt')
    lines = []
    lines.append("ANONYMIZATION MAPPING")
    lines.append("=" * 50)
    lines.append("")
    
    for tag, values in sorted(replacements.items()):
        unique_values = sorted(set(values))
        lines.append(f"{tag}:")
        for value in unique_values:
            lines.append(f"  - {value}")
        lines.append("")
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    print(f"Mapping files saved:")
    print(f"  - {json_path}")
    print(f"  - {txt_path}")

# ========== Main Application ==========

def main():
    """Main application entry point"""
    parser = argparse.ArgumentParser(description='Simple Czech Document Anonymizer')
    parser.add_argument('input', help='Input document path')
    parser.add_argument('--output', '-o', help='Output document path')
    
    args = parser.parse_args()
    
    # Setup logging
    logger = setup_logging()
    
    # Validate input
    input_path = Path(args.input)
    if not input_path.exists():
        logger.error(f"Input file not found: {input_path}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}_anonymized{input_path.suffix}"
    
    try:
        # Initialize anonymizer
        logger.info("Initializing anonymizer...")
        anonymizer = SimpleAnonymizer()
        
        # Initialize processor
        processor = DocumentProcessor(anonymizer)
        
        # Process document
        logger.info(f"Processing document: {input_path}")
        if input_path.suffix.lower() == '.docx':
            result = processor.process_docx(input_path, output_path)
        else:
            result = processor.process_text(input_path, output_path)
        
        # Save mappings
        mapping_base = output_path.parent / f"{output_path.stem}_mapping"
        save_mappings(mapping_base, result.replacements, result.statistics)
        
        # Print results
        logger.info("Anonymization completed successfully!")
        logger.info(f"Output document: {output_path}")
        logger.info(f"Processing time: {result.processing_time:.2f} seconds")
        logger.info(f"Statistics: {result.statistics}")
        
        print(f"\n✅ Anonymization completed!")
        print(f"📄 Output: {output_path}")
        print(f"📊 Statistics: {result.statistics}")
        print(f"⏱️  Time: {result.processing_time:.2f}s")
        
    except Exception as e:
        logger.error(f"Anonymization failed: {e}")
        print(f"❌ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()