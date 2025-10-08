# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer — v5.0
- FIX: Eliminace falešných PERSON z rolí a frází (např. "Pronajímatel Nájemci", "Užívat Byt")
- Přísnější fallback na osoby: vyžaduje "first-name-like" 1. token a zákaz rolových slov
- Rozšířený blacklist slov (role/terminologie smluv)
- Zachováno: kontextová disambiguace OP vs. RČ vs. BANK, sjednocení pádů (Říha/Novotná, -ek/-ec), post-merge, duplicit fix, word-boundary mapy, očista neviditelných znaků
Výstupy: <basename>_anon.docx / _map.json / _map.txt
"""

import sys, re, json, unicodedata
from typing import Optional
from pathlib import Path
from collections import defaultdict, OrderedDict
from docx import Document

# =============== Utility ===============
INVISIBLE = '\u00ad\u200b\u200c\u200d\u2060\ufeff'  # SHY, ZWSP, ZWNJ, ZWJ, WJ, BOM

def clean_invisibles(text: str) -> str:
    if not text: return ''
    text = text.replace('\u00a0', ' ')
    return re.sub('['+re.escape(INVISIBLE)+']', '', text)

def normalize_for_matching(text: str) -> str:
    if not text: return ""
    n = unicodedata.normalize('NFD', text)
    no_diac = ''.join(c for c in n if not unicodedata.combining(c))
    return re.sub(r'[^A-Za-z]', '', no_diac).lower()

def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

def get_text(p) -> str:
    return ''.join(r.text or '' for r in p.runs) or p.text or ''

def set_text(p, s: str):
    if p.runs:
        p.runs[0].text = s
        for r in p.runs[1:]: r.text = ''
    else:
        p.text = s

def preserve_case(surface: str, tag: str) -> str:
    if surface.isupper(): return tag.upper()
    if surface.istitle(): return tag
    return tag

# =============== Lexika (zkrácené core, lze rozšiřovat) ===============
CZECH_FIRST_NAMES = {
    # Mužská (výběr + doplnění problematik)
    "jiří","jan","petr","josef","pavel","martin","jaroslav","tomáš","miroslav","františek",
    "zdeněk","václav","michal","milan","vladimír","jakub","karel","lukáš","ladislav","david",
    "ondřej","stanislav","marek","roman","robert","daniel","radek","aleš","matěj","adam",
    "antonín","filip","ivan","radovan","vojtěch","libor","richard","dalibor","rostislav",
    "vít","přemysl","arnošt","bruno","cyril","dominik","emil","erik","evžen","jaromír",
    "jindřich","julius","konrád","marcel","matouš","maxmilián","miloš","norbert","otakar",
    "patrik","radim","robin","rudolf","samuel","sebastián","šimon","štefan","tadeáš","vilém",
    # Ženská (výběr)
    "marie","jana","eva","hana","anna","lenka","kateřina","lucie","věra","alena",
    "petra","veronika","jaroslava","martina","ivana","zuzana","michaela","jitka","monika","andrea",
    "barbora","kristýna","markéta","tereza","klára","pavla","simona","natálie","ludmila","dagmar",
    "pavlína","radka","adéla","aneta","eliška","soňa","viktorie","alžběta","miriam","nikola",
}

# Termíny, které často vypadají jako příjmení, ale nejsou osoby
SURNAME_BLACKLIST = {
    'smlouva','smlouvě','smlouvy','smlouvou','článek','článku','články',
    'datum','číslo','adresa','bydliště','průkaz','občanský','rodné','zákon','sb','kč','čr',
    'ustanovení','příloha','titul','oddíl','bod','pověřený','zástupce','nájem','pronájem',
    'byt','nájemci','nájemce','pronajímatel','pronajímateli','pronajímateli','pronajímateli,',
    'užívat','hlásit','nepřenechávat','elektřina','plyn','sconto','bolton','předat','předání',
    'cena','kauce','záloha','platba','sankce','odpovědnost','poškození','opravy','závady'
}

# Role slova (tvrdý stop pro osoby)
ROLE_STOP = {
    'pronajímatel','nájemce','dlužník','věřitel','objednatel','zhotovitel',
    'zaměstnanec','zaměstnavatel','ručitel','spoludlužník','jednatel','svědek',
    'statutární','zástupce','pojistník','pojištěný','odesílatel','příjemce'
}

# =============== Inference: nominativ ===============
def _male_genitive_to_nominative(obs: str) -> Optional[str]:
    """Heuristiky pro převod genitivu mužských jmen zpět na nominativ (Radka→Radek, Pavla→Pavel, Marka→Marek)."""
    lo = obs.lower()
    cands = []
    if lo.endswith('ka') and len(obs) > 2:   # Radka -> Radek
        cands.append(obs[:-2] + 'ek')
    if lo.endswith('la') and len(obs) > 2:   # Pavla -> Pavel
        cands.append(obs[:-2] + 'el')
    if lo.endswith('a') and len(obs) > 1:    # Mareka -> Marek (fallback)
        cands.append(obs[:-1])
    for cand in cands:
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
    return None

def infer_first_name_nominative(observed: str, surname_observed: str = "") -> Optional[str]:
    """
    Context-aware inference:
    - Nejprve disambiguace genitivu mužských jmen, pokud příjmení NEvypadá žensky (není -ová/-á/-ou/-é).
    - Až poté přímý match v seznamu.
    - Pak pády (ženské -a; mužské -ovi/-em/-e/-u).
    """
    if not observed: return None
    obs = observed.strip()
    surname_lower = (surname_observed or "").lower()
    female_like_surname = surname_lower.endswith(('ová', 'á', 'ou', 'é'))

    if not female_like_surname:
        cand = _male_genitive_to_nominative(obs)
        if cand: return cand

    norm = normalize_for_matching(obs)
    if norm in CZECH_FIRST_NAMES:
        return obs

    for suf in ['inou','ině','inu','iny','ou','u','y','e','ě','o']:
        if obs.lower().endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)] + 'a'
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    for suf in ['ovi','em','e','u']:
        if obs.lower().endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)]
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand
    return None

def infer_surname_nominative(observed: str) -> str:
    """Nominativ příjmení: -ová, adj. -á/-ý, maskulin -a, -ek/-ec paradigmata, obecné maskulina."""
    if not observed: return observed
    obs = observed.strip()
    low = obs.lower()

    # ženské -ová: ...ovou / ...ové → ...ová
    if low.endswith('ovou') and len(obs) > 4: return obs[:-4] + 'ová'
    if low.endswith('ové') and len(obs) > 3:  return obs[:-3] + 'á'

    # adjektivní ženské: ...é / ...ou → ...á
    if low.endswith('é') and len(obs) > 2:    return obs[:-1] + 'á'
    if low.endswith('ou') and not low.endswith('ovou') and len(obs) > 2:
        return obs[:-2] + 'á'

    # -ek → -k- paradigmata (Mareček -> Marečka/Marečkovi/…)
    m = re.match(r'^(.*)čk(a|ovi|em|u|e|y|ou|ům|ách)?$', obs, flags=re.IGNORECASE)
    if m:
        base = m.group(1)
        return base + 'ček'
    m2 = re.match(r'^(.*)nk(a|ovi|em|u|e|y|ou|ům|ách)?$', obs, flags=re.IGNORECASE)
    if m2:
        base = m2.group(1)
        return base + 'nek'
    if low.endswith(('ka','kovi','kem','ku','ke')) and len(obs) > 3:
        return re.sub(r'k(ovi|em|u|e|a)?$', 'ek', obs, flags=re.IGNORECASE)

    # -ec → -c- paradigmata (Samec -> Samce/Samci/…)
    m3 = re.match(r'^(.*)c(e|i|em|ů|ích|ům|ech|emi|u|y)?$', obs, flags=re.IGNORECASE)
    if m3:
        base = m3.group(1)
        return base + 'ec'

    # maskulin -ovi (Říhovi → Říha)
    if low.endswith('ovi') and len(obs) > 4:  return obs[:-3] + 'a'

    # maskulin -em/-e/-u/-y (Říhou/Říhe/Říhu/Říhy → Říha)
    for suf in ('em','e','u','y'):
        if low.endswith(suf) and len(obs) > len(suf)+1:
            return obs[:-len(suf)] + 'a'

    return obs

# =============== Varianty pro nahrazování ===============
def variants_for_first(first: str) -> set:
    f = first.strip()
    if not f: return {''}
    V = {f, f.lower(), f.capitalize()}
    low = f.lower()
    if low.endswith('a'):
        stem = f[:-1]
        V |= {stem+'y', stem+'e', stem+'ě', stem+'u', stem+'ou', stem+'o'}
        V |= {stem+s for s in ['in','ina','iny','ině','inu','inou','iným','iných']}
        if stem.endswith('tr'):
            V |= {stem[:-1]+'ř'+s for s in ['in','ina','iny','ině','inu','inou','iným','iných']}
    else:
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}
        V |= {f+'ův'} | {f+'ov'+s for s in ['a','o','y','ě','ým','ých']}
        if low.endswith('ek'): V.add(f[:-2] + 'ka')  # Radek→Radka
        if low.endswith('el'): V.add(f[:-2] + 'la')  # Pavel→Pavla
    V |= {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(V)}
    return V

def variants_for_surname(surname: str) -> set:
    s = surname.strip()
    if not s: return {''}
    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    if low.endswith('ová'):
        base = s[:-1]                      # ...ov + á
        out |= {s, base+'é', base+'ou'}    # Svobodová/Svobodové/Svobodovou
        return out
    if low.endswith(('ský','cký','ý')):
        stem = s[:-1] if low.endswith('ý') else s[:-3]
        out |= {stem+'ý', stem+'ého', stem+'ému', stem+'ým', stem+'ém', stem+'á', stem+'é', stem+'ou'}
        return out
    if low.endswith('á'):
        stem = s[:-1]; out |= {s, stem+'é', stem+'ou'}; return out
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {s, stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e', stem_k+'y', stem_k+'ou'}
        return out
    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {s, stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'ů', stem_c+'ům', stem_c+'ích', stem_c+'ech', stem_c+'emi', stem_c+'u', stem_c+'y'}
        return out
    if low.endswith('a') and len(s) >= 2:
        stem = s[:-1]
        out |= {s, stem+'y', stem+'ovi', stem+'ou', stem+'u', stem+'e'}
        return out
    out |= {s+'a', s+'ovi', s+'e', s+'em', s+'u'}
    return out

# =============== Ostatní entity (regexy) ===============
ADDRESS_RE = re.compile(r'(?<!\[)\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][^\n\r,\[\]]{2,50}?\s+\d{1,4}(?:/\d{1,4})?,[ \t]*\d{3}[ \t]?\d{2}[ \t]+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][^\n\r,\[\]]{1,40}\b', re.UNICODE)
ACCT_RE    = re.compile(r'\b(?:\d{1,6}-)?\d{2,10}/\d{4}\b')
BIRTHID_RE = re.compile(r'\b\d{6}\s*/\s*\d{3,4}\b')
IDCARD_RE  = re.compile(r'\b\d{6,9}/\d{3,4}\b|\b\d{9}\b|[A-Z]{2,3}[ \t]?\d{6,9}\b')
PHONE_RE   = re.compile(r'(?<!\d)(?:\+420|00420)?[ \t\-]?\d{3}[ \t\-]?\d{3}[ \t\-]?\d{3}(?!\s*/\d{4})\b')
EMAIL_RE   = re.compile(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}')
DATE_RE    = re.compile(r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b')
STATUTE_RE = re.compile(r'\b(Sb\.?|zákon(a|u)?|zákon\s*č\.)\b', re.IGNORECASE)
PAIR_RE    = re.compile(r'(?<!\w)([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]{1,})\s+([A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]{1,})(?!\w)')
TITLES_RE  = re.compile(r'\b(Mgr|Ing|Dr|Ph\.?D|RNDr|MUDr|JUDr|PhDr|PaedDr|ThDr|RCDr|MVDr|DiS|Bc|BcA|MBA|LL\.?M|prof|doc)\.?\s+', re.IGNORECASE)

# Kontextové klíčové výrazy
CTX_OP     = re.compile(r'\b(OP|občansk(ý|ého|ému|ém|ým)|průkaz|č\.\s*OP)\b', re.IGNORECASE)
CTX_BIRTH  = re.compile(r'\b(rodn[ée]\s*č[íi]slo|RČ|rodn[ée])\b', re.IGNORECASE)
CTX_BANK   = re.compile(r'\b(účet|účtu|účtem|bankovní|veden[eyá].*u|banka|banky|IBAN)\b', re.IGNORECASE)
CTX_PERSON = re.compile(
    r'(nar\.|narozen|rodn[ée]\s*č[íi]slo|RČ|bytem|trval[é]\s*bydlišt[ěi]|'
    r'(e-?mail)|tel\.?|telefon|č\.\s*účtu|IBAN|SPZ|Mgr\.|Ing\.|Bc\.|PhDr\.|JUDr\.)',
    re.IGNORECASE
)
CTX_ROLE   = re.compile(r'\b(pronaj[ií]matel|n[aá]jemce|dlu[zž]n[ií]k|v[eě]řitel|objednatel|zhotovitel|zam[eě]stnanec|zam[eě]stnavatel|ručitel|spoludlu[zž]n[ií]k|jednatel|statut[aá]rn[ií]\s+z[aá]stupce|sv[eě]dek)\b', re.IGNORECASE)
CTX_LABEL  = re.compile(r'j[mn][eě]no\s*(,|a)?\s*př[ií]jmen[ií]', re.IGNORECASE)

# ======== Heuristika: vypadá 1. token jako křestní jméno? ========
def looks_like_firstname(token: str) -> bool:
    if not token or not token[0].isupper(): return False
    norm = normalize_for_matching(token)
    if norm in CZECH_FIRST_NAMES: return True
    # jednoduché morfologické indicie (konzervativní)
    return any([
        norm.endswith('ek'),     # Radek, Marek
        norm.endswith('el'),     # Pavel, Karel
        norm.endswith('oš'),     # Miloš
        norm.endswith('oš'.replace('š','s')), # os fallback ascii
        norm.endswith('áš'),     # Lukáš, Tomáš
        norm.endswith('an'),     # Roman, Ivan
        norm.endswith('en'),     # Jindřich -> ne, ale ponecháme mírné
        norm.endswith('a') and len(norm) > 3,  # ženská na -a
    ])

# =============== Anonymizer ===============
class Anonymizer:
    def __init__(self, verbose=False):
        self.verbose = verbose
        self.counter = defaultdict(int)
        self.tag_map = defaultdict(list)
        self.value_to_tag = {}  # duplicit fix: (cat:value) -> tag
        self.person_index = {}  # (norm_first, norm_last) -> tag
        self.canonical_persons = []   # [{'first','last','tag'}]
        self.person_variants = {}     # tag -> set(variants)
        self.source_text = ""

    def _get_or_create_tag(self, cat: str, value: str) -> str:
        norm_val = ' '.join(value.split())
        lookup_key = f"{cat}:{norm_val}"
        if lookup_key in self.value_to_tag:
            return self.value_to_tag[lookup_key]
        self.counter[cat] += 1
        tag = f'[[{cat}_{self.counter[cat]}]]'
        self.value_to_tag[lookup_key] = tag
        self._record_value(tag, value)
        return tag

    def _record_value(self, tag: str, value: str):
        if value and re.search(r'(?<!\w)'+re.escape(value)+r'(?!\w)', self.source_text):
            if value not in self.tag_map[tag]:
                self.tag_map[tag].append(value)

    def _ensure_person_tag(self, first_nom: str, last_nom: str) -> str:
        key = (normalize_for_matching(first_nom), normalize_for_matching(last_nom))
        if key in self.person_index:
            return self.person_index[key]
        tag = self._get_or_create_tag('PERSON', f'{first_nom} {last_nom}')
        self.person_index[key] = tag
        self.canonical_persons.append({'first': first_nom, 'last': last_nom, 'tag': tag})
        fvars = variants_for_first(first_nom)
        svars = variants_for_surname(last_nom)
        self.person_variants[tag] = {f'{f} {s}' for f in fvars for s in svars}
        return tag

    # --- detekce osob (přísnější fallback) ---
    def _extract_persons_to_index(self, text: str):
        text_no_titles = TITLES_RE.sub('', text)
        for m in PAIR_RE.finditer(text_no_titles):
            s, e = m.span()
            f_tok, l_tok = m.group(1), m.group(2)

            # tvrdý stop na role / běžné termíny
            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            l_nom = infer_surname_nominative(l_tok)

            # 1) whitelisted křestní jméno
            if normalize_for_matching(f_nom) in CZECH_FIRST_NAMES:
                self._ensure_person_tag(f_nom, l_nom)
                continue

            # 2) fallback přes PERSON/ROLE/LABEL kontext + first-name-like + zákaz role slov
            pre = text[max(0, s-160):s]; post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)
            if (has_ctx
                and f_tok[:1].isupper() and l_tok[:1].isupper()
                and looks_like_firstname(f_tok)
                and f_tok.lower() not in ROLE_STOP and l_tok.lower() not in ROLE_STOP):
                self._ensure_person_tag(f_nom, l_nom)

    def _apply_known_people(self, text: str) -> str:
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])
            for pat in sorted(self.person_variants[tag], key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(pat)+r'(?!\w)', re.IGNORECASE)
                def repl(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl, text)
            # přivlastňovací
            first_low, last_low = p['first'].lower(), p['last'].lower()
            poss = set()
            if first_low.endswith('a'):
                stem = p['first'][:-1]
                poss |= {stem+s for s in ['in','ina','iny','ině','inu','inou','iným','iných']}
                if stem.endswith('tr'):
                    poss |= {stem[:-1]+'ř'+s for s in ['in','ina','iny','ině','inu','inou','iným','iných']}
            else:
                poss |= {p['first']+'ův'} | {p['first']+'ov'+s for s in ['a','o','y','ě','ým','ých']}
            if not last_low.endswith('ová'):
                poss |= {p['last']+'ův'} | {p['last']+'ov'+s for s in ['a','o','y','ě','ým','ých']}
            for token in sorted(list(poss), key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(token)+r'(?!\w)', re.IGNORECASE)
                def repl2(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl2, text)
        return text

    def _replace_remaining_people(self, text: str) -> str:
        text_no_titles = TITLES_RE.sub('', text)
        offset = 0
        for m in list(PAIR_RE.finditer(text_no_titles)):
            s, e = m.start()+offset, m.end()+offset
            seg = text[s:e]
            if seg.startswith('[[') and seg.endswith(']]'):
                continue
            f_tok, l_tok = m.group(1), m.group(2)

            # role/blacklist stop
            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            pre = text[max(0, s-160):s]; post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)

            if (normalize_for_matching(f_nom) not in CZECH_FIRST_NAMES
                and not (has_ctx and looks_like_firstname(f_tok))):
                continue

            l_nom = infer_surname_nominative(l_tok)
            tag = self._ensure_person_tag(f_nom, l_nom)
            before = text
            text = text[:s] + preserve_case(seg, tag) + text[e:]
            self._record_value(tag, seg)
            offset += len(text) - len(before)
        return text

    # --- ostatní entity ---
    def _is_statute(self, text: str, s: int, e: int) -> bool:
        pre = text[max(0, s-20):s]; post = text[e:e+10]
        return bool(STATUTE_RE.search(pre) or STATUTE_RE.search(post))

    def _replace_entity(self, text: str, rx: re.Pattern, cat: str) -> str:
        def repl(m):
            v = m.group(0)
            tag = self._get_or_create_tag(cat, v)
            self._record_value(tag, v)
            return tag
        return rx.sub(repl, text)

    def anonymize_entities(self, text: str) -> str:
        # EMAIL
        text = self._replace_entity(text, EMAIL_RE, 'EMAIL')

        # ADRESA
        def addr_repl(m):
            v = m.group(0).strip()
            v = re.sub(r'^(Trvalé\s+bydliště|Bydliště|Adresa)\s*:\s*','', v, flags=re.IGNORECASE)
            tag = self._get_or_create_tag('ADDRESS', v); self._record_value(tag, v); return tag
        text = ADDRESS_RE.sub(addr_repl, text)

        # DATUM
        text = self._replace_entity(text, DATE_RE, 'DATE')

        # TELEFON (vyhnout se OP kontextu)
        def phone_repl(m):
            v = m.group(0); s,e = m.span()
            pre = text[max(0, s-15):s]
            if re.search(r'(OP|občansk\w+|č\.\s*OP)', pre, re.IGNORECASE):
                tag = self._get_or_create_tag('ID_CARD', v); self._record_value(tag, v); return tag
            if re.match(r'^\s*/\d{4}', text[e:e+6]): return v
            tag = self._get_or_create_tag('PHONE', v); self._record_value(tag, v); return tag
        text = PHONE_RE.sub(phone_repl, text)

        # BANK/OP heuristika pro "xx/xxxx" (účty, OP) – rozhoduj kontextem
        def acct_like(m):
            s,e = m.span()
            if self._is_statute(text, s, e): return m.group(0)
            raw = m.group(0)
            pre = text[max(0, s-30):s]; post = text[e:e+30]
            if CTX_BANK.search(pre+post):
                tag = self._get_or_create_tag('BANK', raw); self._record_value(tag, raw); return tag
            if CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', raw); self._record_value(tag, raw); return tag
            return raw
        text = ACCT_RE.sub(acct_like, text)

        # RODNÉ ČÍSLO / OP – čistý formát r.č. → rozhodni kontextem
        def birth_or_id_repl(m):
            v = m.group(0); s,e = m.span()
            pre = text[max(0, s-30):s]; post = text[e:e+30]
            if CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', v)
            elif CTX_BIRTH.search(pre+post):
                tag = self._get_or_create_tag('BIRTH_ID', v)
            else:
                tag = self._get_or_create_tag('BIRTH_ID', v)  # konzervativně r.č.
            self._record_value(tag, v)
            return tag
        text = BIRTHID_RE.sub(birth_or_id_repl, text)

        # OP – jiné formáty (9 číslic, alfanumerické prefixy)
        def id_repl(m):
            v = m.group(0)
            tag = self._get_or_create_tag('ID_CARD', v); self._record_value(tag, v); return tag
        text = IDCARD_RE.sub(id_repl, text)

        return text

    # --- post-merge PERSON tagů ---
    def post_merge_person_tags(self, doc: Document):
        key_to_tags = defaultdict(set)
        for tag, vals in list(self.tag_map.items()):
            if not tag.startswith('[[PERSON_'): continue
            for v in vals:
                m = PAIR_RE.search(v)
                if not m: continue
                f_nom = infer_first_name_nominative(m.group(1), m.group(2)) or m.group(1)
                l_nom = infer_surname_nominative(m.group(2))
                key = (normalize_for_matching(f_nom), normalize_for_matching(l_nom))
                key_to_tags[key].add(tag)

        redirect = {}
        for key, tags in key_to_tags.items():
            if len(tags) <= 1: continue
            canon = sorted(tags)[0]
            for t in tags:
                if t != canon:
                    redirect[t] = canon

        if redirect:
            for p in iter_paragraphs(doc):
                txt = get_text(p)
                new = txt
                for src, dst in redirect.items():
                    new = new.replace(src, dst)
                if new != txt:
                    set_text(p, new)

            for src, dst in redirect.items():
                if src in self.tag_map:
                    for v in self.tag_map[src]:
                        if v not in self.tag_map[dst]: self.tag_map[dst].append(v)
                    del self.tag_map[src]

    # --- hlavní průchod ---
    def anonymize_docx(self, input_path: str, output_path: str, json_map: str, txt_map: str):
        doc = Document(input_path)
        pieces = []
        for p in iter_paragraphs(doc):
            pieces.append(clean_invisibles(get_text(p)))
        self.source_text = '\n'.join(pieces)

        # 1) registrace osob
        self._extract_persons_to_index(self.source_text)

        # 2) průchod: osoby → ostatní entity
        for p in iter_paragraphs(doc):
            raw = get_text(p)
            if not raw.strip(): continue
            txt = clean_invisibles(raw)
            txt = self._apply_known_people(txt)
            txt = self._replace_remaining_people(txt)
            txt = self.anonymize_entities(txt)
            if txt != raw:
                set_text(p, txt)

        # 3) post-merge
        self.post_merge_person_tags(doc)

        # 4) uložení
        doc.save(output_path)

        # 5) mapy
        data = OrderedDict((tag, self.tag_map[tag]) for tag in sorted(self.tag_map.keys()))
        with open(json_map, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        with open(txt_map, 'w', encoding='utf-8') as f:
            sections = [
                ("OSOBY", "PERSON"),
                ("RODNÁ ČÍSLA", "BIRTH_ID"),
                ("BANKOVNÍ ÚČTY", "BANK"),
                ("TELEFONY", "PHONE"),
                ("EMAILY", "EMAIL"),
                ("OBČANSKÉ PRŮKAZY", "ID_CARD"),
                ("DATA", "DATE"),
                ("ADRESY", "ADDRESS"),
            ]
            for title, pref in sections:
                items = []
                for tag, vals in sorted(self.tag_map.items()):
                    if tag.startswith(f'[[{pref}_'):
                        for v in vals:
                            items.append(f"{tag}: {v}")
                if items:
                    f.write(f"{title}\n{'-'*len(title)}\n")
                    f.write("\n".join(items) + "\n\n")

# =============== CLI ===============
def main():
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("docx_path", nargs='?', help="Cesta k .docx souboru")
    args = ap.parse_args()

    path = Path(args.docx_path) if args.docx_path else Path(input("Přetáhni sem .docx soubor nebo napiš cestu: ").strip().strip('"'))
    if not path.exists():
        print("❌ Soubor nenalezen:", path); return 2
    base = path.stem
    out_docx = path.parent / f"{base}_anon.docx"
    out_json = path.parent / f"{base}_map.json"
    out_txt  = path.parent / f"{base}_map.txt"
    a = Anonymizer(verbose=False)
    a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))
    print("✅ Výstupy:")
    print(" -", out_docx)
    print(" -", out_json)
    print(" -", out_txt)

if __name__ == "__main__":
    sys.exit(main())
