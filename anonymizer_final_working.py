# -*- coding: utf-8 -*-
"""
Final Working Czech Document Anonymizer
=======================================
A complete and reliable anonymization system for Czech documents.

Usage:
    python anonymizer_final_working.py input.docx [--output output.docx] [--level full]
"""

import sys
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Set, Tuple, Optional
from dataclasses import dataclass
from enum import Enum
import argparse
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ========== Configuration ==========

class AnonymizationLevel(Enum):
    MINIMAL = "minimal"
    STANDARD = "standard"
    FULL = "full"

@dataclass
class AnonymizationResult:
    original_text: str
    anonymized_text: str
    replacements: Dict[str, List[str]]
    statistics: Dict[str, int]
    processing_time: float

# ========== Utility Functions ==========

def setup_logging(level: str = "INFO") -> logging.Logger:
    """Setup logging configuration"""
    logging.basicConfig(
        level=getattr(logging, level.upper()),
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('anonymizer.log', encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger(__name__)

# ========== Czech Name Detection ==========

class CzechNameDetector:
    """Czech name detection with comprehensive database"""
    
    def __init__(self):
        self.male_names = {
            'jan', 'petr', 'pavel', 'tomáš', 'martin', 'jaroslav', 'milan', 'františek',
            'josef', 'antonín', 'zdeněk', 'vladimír', 'stanislav', 'luděk', 'karel',
            'michal', 'david', 'lukáš', 'ondřej', 'jakub', 'matěj', 'adam', 'daniel',
            'filip', 'mikuláš', 'vít', 'matyáš', 'kryštof', 'sebastian', 'benjamin',
            'ondra', 'honza', 'pepa', 'míra', 'jirka', 'kuba', 'tonda', 'václav',
            'jindřich', 'bohumil', 'bohuslav', 'břetislav', 'ctibor', 'dalibor',
            'dušan', 'emil', 'erik', 'evžen', 'gustav', 'henry', 'hynek', 'igor',
            'ivan', 'kamil', 'lubomír', 'ludvík', 'marcel', 'marek', 'miloslav',
            'mirek', 'oldřich', 'radim', 'radomír', 'rené', 'richard', 'robert',
            'roman', 'rudolf', 'sylvestr', 'tadeáš', 'tibor', 'tobias', 'vlastimil',
            'vratislav', 'zbynek', 'zdislav'
        }
        
        self.female_names = {
            'marie', 'jana', 'eva', 'hana', 'anna', 'věra', 'alena', 'lenka',
            'kateřina', 'lucie', 'petra', 'zuzana', 'iveta', 'monika', 'veronika',
            'tereza', 'barbora', 'adéla', 'karolína', 'kristýna', 'nikola', 'natálie',
            'eliska', 'sophie', 'emma', 'olivia', 'amélie', 'aneta', 'klára', 'julie',
            'alžběta', 'andrea', 'blanka', 'božena', 'dana', 'daniela', 'drahomíra',
            'edita', 'emilie', 'gabriela', 'helena', 'ivana', 'jarmila', 'jindřiška',
            'jolana', 'josefa', 'květa', 'ladislava', 'libuše', 'ludmila', 'magdalena',
            'marcela', 'markéta', 'milena', 'miriam', 'naděžda', 'olga', 'pavla',
            'radka', 'renata', 'romana', 'sandra', 'simona', 'soňa', 'stanislava',
            'světlana', 'šárka', 'táňa', 'valerie', 'vlasta', 'zdeňka', 'žofie'
        }
        
        self.surnames = {
            'novák', 'svoboda', 'novotný', 'dvořák', 'černý', 'procházka',
            'kučera', 'veselý', 'horák', 'němec', 'pokorný', 'pospíšil',
            'havel', 'bláha', 'krejčí', 'stárek', 'kříž', 'beneš', 'fiala',
            'moravec', 'barták', 'urban', 'polák', 'doležal', 'šimánek',
            'nováková', 'svobodová', 'novotná', 'dvořáková', 'černá', 'procházková',
            'kučerová', 'veselá', 'horáková', 'němcová', 'pokorná', 'pospíšilová',
            'havlová', 'bláhová', 'krejčová', 'stárková', 'křížová', 'benešová',
            'fialová', 'moravcová', 'bartáková', 'urbanová', 'poláková', 'šimánková'
        }
        
        self.surname_suffixes = {'ová', 'ek', 'ík', 'ák', 'ček', 'čík', 'ko', 'ka', 'ja', 'ský', 'cký'}
    
    def is_likely_first_name(self, word: str) -> bool:
        """Check if word is likely a first name"""
        normalized = word.lower().strip()
        return normalized in self.male_names or normalized in self.female_names
    
    def is_likely_surname(self, word: str) -> bool:
        """Check if word is likely a surname"""
        normalized = word.lower().strip()
        return (normalized in self.surnames or 
                any(normalized.endswith(suffix) for suffix in self.surname_suffixes))

# ========== Pattern Detection ==========

class PatternDetector:
    """Pattern detection for sensitive data"""
    
    def __init__(self, level: AnonymizationLevel = AnonymizationLevel.STANDARD):
        self.level = level
        self.patterns = self._initialize_patterns()
    
    def _initialize_patterns(self) -> List[Tuple[str, str, bool]]:
        """Initialize detection patterns"""
        patterns = []
        
        # Date patterns
        patterns.append((r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b', 'DATE', False))
        
        # Czech birth number (RČ)
        patterns.append((r'\b\d{2}[0156]\d{3,4}/\d{4}\b', 'BIRTH_ID', False))
        
        # ID card number (with context)
        patterns.append((r'\b\d{9}\b', 'ID_CARD', True))
        
        # Bank account numbers
        patterns.append((r'\b\d{1,6}-\d{1,10}/\d{4}\b', 'BANK', False))
        patterns.append((r'\bCZ\d{2}(?:\s?\d){20}\b', 'BANK', False))
        
        # Phone numbers
        patterns.append((r'(?:\+?420[\s\-]?)?(?<!\d)(?:\d{3}[\s\-]?){2}\d{3}(?!\d)', 'PHONE', False))
        
        # Email addresses
        patterns.append((r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', 'EMAIL', False))
        
        # Addresses
        patterns.append((r'\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]{2,30}\s+\d{1,4}(?:/\d{1,4})?,\s*\d{3}\s?\d{2}\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž\s]{2,20}\b', 'ADDRESS', False))
        
        if self.level == AnonymizationLevel.FULL:
            patterns.extend([
                (r'\b\d{3}\s?\d{2}\s?\d{3}\b', 'SOCIAL_SECURITY', False),
                (r'\b[A-Z]{2}\d{6}\b', 'PASSPORT', False),
                (r'\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b', 'CREDIT_CARD', False),
                (r'\b(?![IOQ])[A-HJ-NPR-Z0-9]{17}\b', 'VIN', False),
                (r'\b[A-Z]{1,3}\s?[0-9]{1,4}[A-Z]?\b', 'PLATE', False)
            ])
        
        return patterns
    
    def detect_patterns(self, text: str) -> List[Tuple[str, str, int, int]]:
        """Detect all patterns in text"""
        results = []
        
        for pattern, category, needs_context in self.patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                # Check context requirement
                if needs_context:
                    context = text[max(0, match.start()-20):match.end()+20].lower()
                    if not self._has_relevant_context(context, category):
                        continue
                
                # Skip legal references
                if self._is_legal_reference(text, match.start(), match.end()):
                    continue
                
                results.append((category, match.group(), match.start(), match.end()))
        
        return results
    
    def _has_relevant_context(self, context: str, category: str) -> bool:
        """Check if context contains relevant keywords"""
        context_keywords = {
            'id_card': ['op', 'občansk', 'průkaz'],
            'bank_account': ['účet', 'bank', 'čísla'],
            'phone': ['telefon', 'mobil', 'kontakt'],
            'address': ['adresa', 'bydliště', 'ulice']
        }
        
        keywords = context_keywords.get(category, [])
        return any(keyword in context for keyword in keywords)
    
    def _is_legal_reference(self, text: str, start: int, end: int) -> bool:
        """Check if match looks like a legal reference"""
        context = text[max(0, start-15):end+15].lower()
        legal_indicators = ['§', 'zákon', 'oz', 'vyhláška', 'nařízení']
        return any(indicator in context for indicator in legal_indicators)

# ========== Main Anonymizer Class ==========

class WorkingAnonymizer:
    """Working anonymizer with comprehensive detection"""
    
    def __init__(self, level: AnonymizationLevel = AnonymizationLevel.STANDARD):
        self.level = level
        self.name_detector = CzechNameDetector()
        self.pattern_detector = PatternDetector(level)
        self.logger = logging.getLogger(__name__)
        
        # Mapping and tracking
        self.replacements: Dict[str, List[str]] = {}
        self.counters: Dict[str, int] = {}
        self.person_mappings: Dict[Tuple[str, str], str] = {}
    
    def _new_tag(self, category: str) -> str:
        """Generate new anonymization tag"""
        self.counters[category] = self.counters.get(category, 0) + 1
        return f"[[{category}_{self.counters[category]}]]"
    
    def _add_replacement(self, tag: str, original: str) -> None:
        """Add replacement to mapping"""
        if tag not in self.replacements:
            self.replacements[tag] = []
        if original not in self.replacements[tag]:
            self.replacements[tag].append(original)
    
    def _replace_text(self, text: str, start: int, end: int, tag: str, original: str) -> str:
        """Replace text span with anonymization tag"""
        if start < 0 or end > len(text) or start >= end:
            return text
        
        # Check if already anonymized
        if text[start:end].startswith("[[") and text[start:end].endswith("]]"):
            return text
        
        self._add_replacement(tag, original)
        return text[:start] + tag + text[end:]
    
    def anonymize_names(self, text: str) -> str:
        """Anonymize names using improved detection"""
        # Find potential names (capitalized words)
        name_pattern = re.compile(r'\b[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][a-záčďéěíňóřšťúůýž]+\b')
        
        # First pass: find all potential names and their positions
        potential_names = []
        for match in name_pattern.finditer(text):
            word = match.group()
            cleaned = re.sub(r'[^\w]', '', word)
            if (self.name_detector.is_likely_first_name(cleaned) or 
                self.name_detector.is_likely_surname(cleaned)):
                potential_names.append((match.start(), match.end(), word, cleaned))
        
        # Second pass: group names into pairs and singles
        processed_ranges = []
        result_text = text
        
        i = 0
        while i < len(potential_names):
            start, end, word, cleaned = potential_names[i]
            
            # Skip if already processed
            if any(not (end <= existing_start or start >= existing_end) 
                   for existing_start, existing_end in processed_ranges):
                i += 1
                continue
            
            # Look for name pair (first name + surname)
            if (i + 1 < len(potential_names) and 
                self.name_detector.is_likely_first_name(cleaned) and
                self.name_detector.is_likely_surname(potential_names[i + 1][3])):
                
                # Found name pair
                next_start, next_end, next_word, next_cleaned = potential_names[i + 1]
                key = (cleaned.lower(), next_cleaned.lower())
                
                if key not in self.person_mappings:
                    tag = self._new_tag("PERSON")
                    self.person_mappings[key] = tag
                
                tag = self.person_mappings[key]
                self._add_replacement(tag, f"{word} {next_word}")
                
                # Replace both words
                result_text = self._replace_text(result_text, start, next_end, tag, f"{word} {next_word}")
                processed_ranges.append((start, next_end))
                i += 2  # Skip both words
            else:
                # Single name - only if it's a surname or very likely first name
                if (self.name_detector.is_likely_surname(cleaned) or 
                    (self.name_detector.is_likely_first_name(cleaned) and 
                     not any(self.name_detector.is_likely_surname(potential_names[j][3]) 
                            for j in range(i+1, min(i+3, len(potential_names)))))):
                    
                    key = (cleaned.lower(), "")
                    if key not in self.person_mappings:
                        tag = self._new_tag("PERSON")
                        self.person_mappings[key] = tag
                    
                    tag = self.person_mappings[key]
                    self._add_replacement(tag, word)
                    result_text = self._replace_text(result_text, start, end, tag, word)
                    processed_ranges.append((start, end))
                
                i += 1
        
        return result_text
    
    def anonymize_patterns(self, text: str) -> str:
        """Anonymize detected patterns"""
        patterns = self.pattern_detector.detect_patterns(text)
        
        # Sort by position (reverse order to avoid offset issues)
        patterns.sort(key=lambda x: x[2], reverse=True)
        
        # Track processed ranges to avoid conflicts
        processed_ranges = []
        
        for category, match, start, end in patterns:
            # Check if this range overlaps with already processed ranges
            overlaps = any(not (end <= existing_start or start >= existing_end) 
                          for existing_start, existing_end in processed_ranges)
            
            if not overlaps:
                tag = self._new_tag(category)
                text = self._replace_text(text, start, end, tag, match)
                processed_ranges.append((start, end))
        
        return text
    
    def anonymize_text(self, text: str) -> str:
        """Main anonymization method"""
        if not text.strip():
            return text
        
        # Anonymize names first
        text = self.anonymize_names(text)
        
        # Anonymize patterns
        text = self.anonymize_patterns(text)
        
        return text
    
    def get_statistics(self) -> Dict[str, int]:
        """Get anonymization statistics"""
        return dict(self.counters)

# ========== Document Processing ==========

class DocumentProcessor:
    """Process various document formats"""
    
    def __init__(self, anonymizer: WorkingAnonymizer):
        self.anonymizer = anonymizer
        self.logger = logging.getLogger(__name__)
    
    def process_docx(self, input_path: Path, output_path: Path) -> AnonymizationResult:
        """Process DOCX document"""
        start_time = datetime.now()
        
        try:
            doc = Document(str(input_path))
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original = paragraph.text
                    anonymized = self.anonymizer.anonymize_text(original)
                    paragraph.text = anonymized
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original = paragraph.text
                                anonymized = self.anonymizer.anonymize_text(original)
                                paragraph.text = anonymized
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return AnonymizationResult(
                original_text="",
                anonymized_text="",
                replacements=self.anonymizer.replacements,
                statistics=self.anonymizer.get_statistics(),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
    
    def process_text(self, input_path: Path, output_path: Path) -> AnonymizationResult:
        """Process plain text document"""
        start_time = datetime.now()
        
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                original_text = f.read()
            
            anonymized_text = self.anonymizer.anonymize_text(original_text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(anonymized_text)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return AnonymizationResult(
                original_text=original_text,
                anonymized_text=anonymized_text,
                replacements=self.anonymizer.replacements,
                statistics=self.anonymizer.get_statistics(),
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error(f"Error processing text file: {e}")
            raise

# ========== Output and Mapping ==========

def save_mappings(base_path: Path, replacements: Dict[str, List[str]], statistics: Dict[str, int]):
    """Save anonymization mappings in multiple formats"""
    
    # JSON format
    json_path = base_path.with_suffix('.json')
    mapping_data = {
        'metadata': {
            'created_at': datetime.now().isoformat(),
            'statistics': statistics,
            'total_replacements': sum(len(vals) for vals in replacements.values())
        },
        'replacements': replacements
    }
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(mapping_data, f, ensure_ascii=False, indent=2)
    
    # Text format
    txt_path = base_path.with_suffix('.txt')
    lines = []
    lines.append("ANONYMIZACE DOKUMENTU - MAPA NÁHRAD")
    lines.append("=" * 50)
    lines.append("")
    lines.append(f"Vytvořeno: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    lines.append(f"Celkem náhrad: {sum(len(vals) for vals in replacements.values())}")
    lines.append("")
    
    for tag, values in sorted(replacements.items()):
        unique_values = sorted(set(values))
        lines.append(f"{tag}:")
        for value in unique_values:
            lines.append(f"  - {value}")
        lines.append("")
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    # CSV format
    csv_path = base_path.with_suffix('.csv')
    import csv
    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Tag', 'Původní_Hodnota', 'Kategorie'])
        
        for tag, values in sorted(replacements.items()):
            category = tag.split('_')[0].replace('[[', '').replace(']]', '')
            for value in values:
                writer.writerow([tag, value, category])
    
    print(f"Mapovací soubory uloženy:")
    print(f"  - {json_path}")
    print(f"  - {txt_path}")
    print(f"  - {csv_path}")

# ========== Main Application ==========

def main():
    """Main application entry point"""
    parser = argparse.ArgumentParser(
        description='Final Working Czech Document Anonymizer',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python anonymizer_final_working.py document.docx
  python anonymizer_final_working.py document.txt --level full --output anonymized.txt
  python anonymizer_final_working.py data.docx --level standard --log-level DEBUG
        """
    )
    
    parser.add_argument('input', help='Input document path')
    parser.add_argument('--output', '-o', help='Output document path')
    parser.add_argument('--level', choices=['minimal', 'standard', 'full'], 
                       default='standard', help='Anonymization level (default: standard)')
    parser.add_argument('--log-level', choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'],
                       default='INFO', help='Logging level (default: INFO)')
    
    args = parser.parse_args()
    
    # Setup logging
    logger = setup_logging(args.log_level)
    
    # Validate input
    input_path = Path(args.input)
    if not input_path.exists():
        logger.error(f"Input file not found: {input_path}")
        print(f"❌ Chyba: Soubor nenalezen: {input_path}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}_anonymized{input_path.suffix}"
    
    # Determine anonymization level
    level = AnonymizationLevel(args.level)
    
    try:
        # Initialize anonymizer
        logger.info(f"Initializing anonymizer with level: {level.value}")
        print(f"🔧 Inicializace anonymizátoru s úrovní: {level.value}")
        
        anonymizer = WorkingAnonymizer(level)
        
        # Initialize processor
        processor = DocumentProcessor(anonymizer)
        
        # Process document
        logger.info(f"Processing document: {input_path}")
        print(f"📄 Zpracovávám dokument: {input_path}")
        
        if input_path.suffix.lower() == '.docx':
            result = processor.process_docx(input_path, output_path)
        else:
            result = processor.process_text(input_path, output_path)
        
        # Save mappings
        mapping_base = output_path.parent / f"{output_path.stem}_mapping"
        save_mappings(mapping_base, result.replacements, result.statistics)
        
        # Print results
        logger.info("Anonymization completed successfully!")
        logger.info(f"Output document: {output_path}")
        logger.info(f"Processing time: {result.processing_time:.2f} seconds")
        logger.info(f"Statistics: {result.statistics}")
        
        print(f"\n✅ Anonymizace dokončena úspěšně!")
        print(f"📄 Výstupní dokument: {output_path}")
        print(f"📊 Statistiky: {result.statistics}")
        print(f"⏱️  Čas zpracování: {result.processing_time:.2f}s")
        
        # Show preview of anonymized content
        if result.anonymized_text:
            print(f"\n📋 Náhled anonymizovaného obsahu:")
            preview = result.anonymized_text[:300] + "..." if len(result.anonymized_text) > 300 else result.anonymized_text
            print(preview)
        
    except Exception as e:
        logger.error(f"Anonymization failed: {e}")
        print(f"❌ Chyba: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()