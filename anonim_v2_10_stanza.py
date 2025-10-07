# -*- coding: utf-8 -*-
"""
CZ Anonymizer – PRO v1 (lemma + heuristics, 2-pass)
--------------------------------------------------
- Stanza (cs) tokenize + POS + lemma (NO NER needed)
- Pass 1: detect FirstName + Surname pairs (tolerates punctuation/conjunction between)
- Pass 2: replace single tokens that match known first/last lemmas (+ possessives, -ová mapping)
- Extra detectors: birth dates, RČ, OP, bank accounts (prefix-acc/bank, IBAN CZ), phone, email,
  VIN, license plate (SPZ/RZ), addresses (basic pattern).
- Works on DOCX (paragraphs + tables). Saves mapping (JSON + TXT).

Install:
  pip install python-docx stanza
Download models once:
  python -c "import stanza,os; os.makedirs('data/models/stanza_cs', exist_ok=True); stanza.download('cs', model_dir='data/models/stanza_cs')"

Optional first-name library:
  data/lib/cz_names.v1.json  (format from earlier step)

Run:
  python anonymizer_cz_pro.py "cesta\\k\\souboru.docx"
"""

import sys, re, json, unicodedata
from pathlib import Path
from typing import List, Tuple, Dict, Set
from docx import Document

# ========== utils ==========

def nfc_lower(s: str) -> str:
    return unicodedata.normalize("NFC", s or "").lower()

def ensure_dirs(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)

# ========== names library (optional) ==========

def find_lib_json() -> Path | None:
    candidates = [
        Path("data/lib/cz_names.v1.json"),
        Path(__file__).resolve().parent / "data" / "lib" / "cz_names.v1.json",
        Path.cwd() / "data" / "lib" / "cz_names.v1.json",
        Path(__file__).resolve().parent / "cz_names.v1.json",
        Path("/mnt/data/cz_names.v1.json"),
    ]
    for c in candidates:
        if c.exists():
            return c
    return None

def load_firstnames(path: Path) -> Dict[str, Set[str]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    m = {nfc_lower(x) for x in data["firstnames"]["M"]}
    f = {nfc_lower(x) for x in data["firstnames"]["F"]}
    return {"M": m, "F": f, "ALL": m | f}

# ========== Stanza (NO NER) ==========

class StanzaPipe:
    def __init__(self, model_dir: str = "data/models/stanza_cs") -> None:
        import stanza
        self.stanza = stanza
        self.nlp = self.stanza.Pipeline(
            lang="cs",
            processors="tokenize,mwt,pos,lemma",
            model_dir=model_dir,
            download_method=None,
            use_gpu=False
        )
    def analyze(self, text: str):
        return self.nlp(text)

# ========== heuristics ==========

PUNCT_LIKE = {"PUNCT", "SYM"}
SKIP_UPOS_IN_BETWEEN = PUNCT_LIKE | {"ADP", "CCONJ", "SCONJ", "PART"}

# Rough surname suffix list for Czech (signal, not rule)
SURNAME_SUFFIXES = tuple([
    "ová","ek","ík","ák","ček","čík","ko","ka","ja","as","es","is","os","us",
    "ský","cký","r","l","n","m","s","z","č","ř","ť","ď","c"
])

def is_likely_surname(lemma: str) -> bool:
    ll = nfc_lower(lemma)
    return ll.endswith(SURNAME_SUFFIXES)

def map_possessive_to_base(lemma: str) -> List[str]:
    out = []
    ll = nfc_lower(lemma)
    if ll.endswith("ův"):
        base = ll[:-2]
        out.append(base)
        if base.endswith("rl"):
            out.append(base + "el")  # karlův -> karel
    if ll.endswith("in"):
        stem = ll[:-2]
        out.append(stem + "a")  # petřin -> petra
        out.append(stem + "e")  # mariin -> marie
    for suf in ("ova","ovo","ovi","ovy","ových","ovou","ové","ov"):
        if ll.endswith(suf):
            stem = ll[:-len(suf)]
            out.extend([stem, stem + "ová"])
            break
    # unique
    seen = set(); res = []
    for x in out:
        if x not in seen:
            seen.add(x); res.append(x)
    return res

def masculine_feminine_variants(last_lemma: str) -> List[str]:
    ll = nfc_lower(last_lemma)
    if ll.endswith("ová"):
        return [ll, ll[:-3]]  # Nováková -> Novák
    return [ll, ll + "ová"]   # Novák -> Nováková

# ========== core class ==========

class AnonymizerPRO:
    def __init__(self, firstnames: Dict[str, Set[str]] | None = None):
        self.firstnames = firstnames or {"ALL": set()}
        self.tag_counter = 1
        self.map_pair_to_tag: Dict[Tuple[str,str], str] = {}
        self.map_first_to_tag: Dict[str, str] = {}
        self.map_last_to_tag: Dict[str, str] = {}
        self.replacements: Dict[str, List[str]] = {}
        self.counters: Dict[str, int] = {}

    # ---- tagging helpers ----
    def _new_tag(self, category: str) -> str:
        self.counters[category] = self.counters.get(category, 0) + 1
        return f"[[{category}_{self.counters[category]}]]"

    def _add_map(self, tag: str, original: str) -> None:
        self.replacements.setdefault(tag, []).append(original)

    def _replace_span(self, text: str, s: int, e: int, tag: str, original: str) -> str:
        if s < 0 or e > len(text) or s >= e:
            return text
        if text[s:e].startswith("[[") and text[s:e].endswith("]]"):
            return text
        self._add_map(tag, original)
        return text[:s] + tag + text[e:]

    # ---- people detection ----
    def _tag_for_person(self, first_lemma: str, last_lemma: str) -> str:
        key = (nfc_lower(first_lemma), nfc_lower(last_lemma))
        if key not in self.map_pair_to_tag:
            tag = self._new_tag("PERSON")
            self.map_pair_to_tag[key] = tag
            self.map_first_to_tag.setdefault(nfc_lower(first_lemma), tag)
            for v in masculine_feminine_variants(last_lemma):
                self.map_last_to_tag.setdefault(v, tag)
        return self.map_pair_to_tag[key]

    def anonymize_people(self, text: str, pipe: 'StanzaPipe') -> str:
        if not text.strip():
            return text

        doc = pipe.analyze(text)
        # Pass 1: pairs
        pairs = []
        for sent in doc.sentences:
            words = sent.words
            # context boosts
            sent_text = sent.text.lower()
            context_boost = any(x in sent_text for x in [
                "nar.", "r.č.", "rodné číslo", "jmén", "jméno a příjmení",
                "bytem", "trvale bytem", "datum narození", "podpis"
            ])
            i = 0
            while i < len(words):
                w = words[i]
                if w.upos != "PROPN":
                    i += 1; continue
                l1 = nfc_lower(w.lemma or w.text)
                is_firstname = (l1 in self.firstnames.get("ALL", set()))
                # look ahead
                j = i + 1
                while j < len(words) and (words[j].upos in SKIP_UPOS_IN_BETWEEN):
                    j += 1
                if j < len(words) and words[j].upos == "PROPN":
                    w2 = words[j]
                    l2 = nfc_lower(w2.lemma or w2.text)
                    likely = is_firstname or context_boost or is_likely_surname(l2) or l2.endswith("ová")
                    if likely:
                        try:
                            s = int(w.start_char); e = int(w2.end_char)
                            pairs.append((s,e,w.text,w2.text,w.lemma or w.text,w2.lemma or w2.text))
                            i = j + 1; continue
                        except Exception:
                            pass
                i += 1

        # replace from end
        for (s,e,fs,ls,fl,ll) in sorted(pairs, key=lambda x: -x[0]):
            tag = self._tag_for_person(fl, ll)
            text = self._replace_span(text, s, e, tag, f"{fs} {ls}")

        # Pass 2: singles using known sets + possessives
        doc = pipe.analyze(text)
        taken: List[Tuple[int,int]] = []
        def free(seg):
            s,e = seg
            for s2,e2 in taken:
                if not (e <= s2 or s >= e2):
                    return False
            return True

        for sent in doc.sentences:
            for w in sent.words:
                if w.upos not in {"PROPN","ADJ"}:
                    continue
                try:
                    s = int(w.start_char); e = int(w.end_char)
                except Exception:
                    continue
                surf = text[s:e]
                if surf.startswith("[[") and surf.endswith("]]"):
                    continue
                lem = nfc_lower(w.lemma or w.text)

                if lem in self.map_first_to_tag:
                    tag = self.map_first_to_tag[lem]
                    if free((s,e)):
                        text = self._replace_span(text, s, e, tag, surf); taken.append((s,e)); continue
                if lem in self.map_last_to_tag:
                    tag = self.map_last_to_tag[lem]
                    if free((s,e)):
                        text = self._replace_span(text, s, e, tag, surf); taken.append((s,e)); continue
                for b in map_possessive_to_base(lem):
                    if b in self.map_first_to_tag:
                        tag = self.map_first_to_tag[b]
                        if free((s,e)):
                            text = self._replace_span(text, s, e, tag, surf); taken.append((s,e)); break
                    if b in self.map_last_to_tag:
                        tag = self.map_last_to_tag[b]
                        if free((s,e)):
                            text = self._replace_span(text, s, e, tag, surf); taken.append((s,e)); break

        return text

    # ---- regex detectors ----
    DATE_RE = re.compile(r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b')
    # Czech birth number (RČ): 9-10 digits, optional slash, basic validation
    RC_RE = re.compile(r'\b\d{2}[0156]\d{3,4}/?\d{4}\b')
    OP9_RE = re.compile(r'\b\d{9}\b')
    BANK_RE = re.compile(r'\b(?:\d{1,6}-)?\d{1,10}/\d{4}\b')  # prefix-main/bank
    IBAN_CZ_RE = re.compile(r'\bCZ\d{2}(?:\s?\d){20}\b', re.IGNORECASE)
    PHONE_RE = re.compile(r'(?:\+?420[\s\-]?)?(?<!\d)(?:\d{3}[\s\-]?){2}\d{3}(?!\d)')
    EMAIL_RE = re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')
    VIN_RE = re.compile(r'\b(?![IOQ])[A-HJ-NPR-Z0-9]{17}\b')
    RZ_RE = re.compile(r'\b[0-9A-Z]{1,3}\s?[0-9]{1,4}\b')  # simple plate
    ADDRESS_RE = re.compile(
        r'(?:Trvalé bydliště:\s*)?' +
        r'[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][^,\n]{2,40}\s+\d{1,4}(?:/\d{1,4})?,\s*' +
        r'\d{3}\s?\d{2}\s+[A-ZÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ][^,\n]{2,30}'
    )

    def _valid_rc(self, s: str) -> bool:
        ss = s.replace("/", "")
        if len(ss) not in (9, 10) or not ss.isdigit():
            return False
        if len(ss) == 10:
            try:
                return int(ss) % 11 == 0 or int(ss) % 11 == 10
            except Exception:
                return False
        return True

    def anonymize_misc(self, text: str) -> str:
        # Date
        for m in list(self.DATE_RE.finditer(text)):
            tag = self._new_tag("DATE"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # RČ (skip if looks like law 89/2012 etc.)
        for m in list(self.RC_RE.finditer(text)):
            if "§" in text[max(0,m.start()-15):m.end()+15]:
                continue
            if self._valid_rc(m.group(0)):
                tag = self._new_tag("BIRTH_ID"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # OP 9 digits with context 'OP'/'občansk'
        for m in list(self.OP9_RE.finditer(text)):
            ctx = text[max(0, m.start()-20): m.end()+20].lower()
            if "op" in ctx or "občansk" in ctx:
                tag = self._new_tag("ID_CARD"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # Bank account
        for m in list(self.BANK_RE.finditer(text)):
            # avoid statutes like 89/2012
            num = m.group(0)
            if re.match(r'^\d{1,3}/\d{4}$', num):
                ctx = text[max(0, m.start()-30): m.end()+30].lower()
                if "zákon" in ctx or "oz" in ctx:
                    continue
            tag = self._new_tag("BANK"); text = self._replace_span(text, m.start(), m.end(), tag, num)

        # IBAN CZ
        for m in list(self.IBAN_CZ_RE.finditer(text)):
            tag = self._new_tag("BANK"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # Phone
        for m in list(self.PHONE_RE.finditer(text)):
            tag = self._new_tag("PHONE"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # Email
        for m in list(self.EMAIL_RE.finditer(text)):
            tag = self._new_tag("EMAIL"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # VIN
        for m in list(self.VIN_RE.finditer(text)):
            tag = self._new_tag("VIN"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # RZ / SPZ (simple)
        for m in list(self.RZ_RE.finditer(text)):
            tag = self._new_tag("PLATE"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        # Address
        for m in list(self.ADDRESS_RE.finditer(text)):
            tag = self._new_tag("ADDRESS"); text = self._replace_span(text, m.start(), m.end(), tag, m.group(0))

        return text

    # ---- pipeline for one string ----
    def anonymize_text(self, text: str, pipe: 'StanzaPipe') -> str:
        # people
        text = self.anonymize_people(text, pipe)
        # misc numbers/ids
        text = self.anonymize_misc(text)
        return text

# ========== DOCX I/O ==========

def read_docx_paras(p: Path) -> List[str]:
    doc = Document(str(p))
    return [par.text for par in doc.paragraphs]

def write_docx_paras(inp: Path, outp: Path, texts: List[str]) -> None:
    doc = Document(str(inp))
    for i, par in enumerate(doc.paragraphs):
        if i < len(texts):
            par.text = texts[i]
    ensure_dirs(outp)
    doc.save(str(outp))

def process_tables(doc, anon: AnonymizerPRO, pipe: 'StanzaPipe'):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        p.text = anon.anonymize_text(p.text, pipe)

def save_maps(base: Path, mapping: Dict[str, List[str]]):
    base.with_suffix(".json").write_text(json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = []
    for tag, vals in mapping.items():
        uniq = sorted(set(vals))
        lines.append(f"{tag}: " + ", ".join(uniq))
    base.with_suffix(".txt").write_text("\n".join(lines), encoding="utf-8")

# ========== CLI ==========

def main():
    if len(sys.argv) < 2:
        print("Použití: python anonymizer_cz_pro.py \"cesta\\k\\souboru.docx\"")
        sys.exit(1)

    docx_in = Path(sys.argv[1])
    if not docx_in.exists():
        print(f"Soubor nenalezen: {docx_in}")
        sys.exit(1)

    # Load firstnames if available
    first = {}
    lib = find_lib_json()
    if lib:
        first = load_firstnames(lib)

    pipe = StanzaPipe(model_dir="data/models/stanza_cs")
    anon = AnonymizerPRO(firstnames=first)

    # paragraphs
    paras = read_docx_paras(docx_in)
    out_paras = [anon.anonymize_text(t, pipe) for t in paras]

    # write paragraphs, then process tables in-place
    out_docx = docx_in.with_name(docx_in.stem + "_anon.docx")
    write_docx_paras(docx_in, out_docx, out_paras)
    # reopen for tables
    doc = Document(str(out_docx))
    process_tables(doc, anon, pipe)
    doc.save(str(out_docx))

    save_maps(docx_in.with_name(docx_in.stem + "_map"), anon.replacements)

    # stats
    print("Stats:", {k: v for k, v in sorted(anon.counters.items())})
    print(f"OK → {out_docx}")

if __name__ == "__main__":
    main()
