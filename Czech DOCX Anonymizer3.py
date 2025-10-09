# -*- coding: utf-8 -*-
"""
Czech DOCX Anonymizer ‚Äì v6.1
- Naƒç√≠t√° jm√©na z JSON knihovny (cz_names.v1.json)
- Opraveno: BANK vs OP, fale≈°n√© osoby, adresy
V√Ωstupy: <basename>_anon.docx / _map.json / _map.txt
"""

import sys, re, json, unicodedata
from typing import Optional, Set
from pathlib import Path
from collections import defaultdict, OrderedDict
from docx import Document

# =============== Utility ===============
INVISIBLE = '\u00ad\u200b\u200c\u200d\u2060\ufeff'

def clean_invisibles(text: str) -> str:
    if not text: return ''
    text = text.replace('\u00a0', ' ')
    return re.sub('['+re.escape(INVISIBLE)+']', '', text)

def normalize_for_matching(text: str) -> str:
    if not text: return ""
    n = unicodedata.normalize('NFD', text)
    no_diac = ''.join(c for c in n if not unicodedata.combining(c))
    return re.sub(r'[^A-Za-z]', '', no_diac).lower()

def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

def get_text(p) -> str:
    return ''.join(r.text or '' for r in p.runs) or p.text or ''

def set_text(p, s: str):
    if p.runs:
        p.runs[0].text = s
        for r in p.runs[1:]: r.text = ''
    else:
        p.text = s

def preserve_case(surface: str, tag: str) -> str:
    if surface.isupper(): return tag.upper()
    if surface.istitle(): return tag
    return tag

# =============== Naƒçten√≠ knihovny jmen ===============
def load_names_library(json_path: str = "cz_names.v1.json") -> Set[str]:
    try:
        script_dir = Path(__file__).parent if '__file__' in globals() else Path.cwd()
        json_file = script_dir / json_path
        
        if not json_file.exists():
            print(f"‚ö†Ô∏è  Varov√°n√≠: {json_path} nenalezen, pou≈æ√≠v√°m pr√°zdnou knihovnu!")
            return set()
        
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        names = set()
        if 'firstnames_no_diac' in data:
            names.update(data['firstnames_no_diac'].get('M', []))
            names.update(data['firstnames_no_diac'].get('F', []))
        
        print(f"‚úì Naƒçteno {len(names)} jmen z knihovny")
        return names
        
    except Exception as e:
        print(f"‚ö†Ô∏è  Chyba p≈ôi naƒç√≠t√°n√≠: {e}")
        return set()

CZECH_FIRST_NAMES = load_names_library()

# =============== Blacklisty ===============
SURNAME_BLACKLIST = {
    'smlouva','smlouvƒõ','smlouvy','smlouvou','ƒçl√°nek','ƒçl√°nku','ƒçl√°nky',
    'datum','ƒç√≠slo','adresa','bydli≈°tƒõ','pr≈Økaz','obƒçansk√Ω','rodn√©','z√°kon','sb','kƒç','ƒçr',
    'ustanoven√≠','p≈ô√≠loha','titul','odd√≠l','bod','povƒõ≈ôen√Ω','z√°stupce','n√°jem','pron√°jem',
    'byt','n√°jemci','n√°jemce','pronaj√≠matel','pronaj√≠mateli',
    'u≈æ√≠vat','hl√°sit','nep≈ôenech√°vat','elekt≈ôina','plyn','sconto','bolton','p≈ôedat','p≈ôed√°n√≠',
    'cena','kauce','z√°loha','platba','sankce','odpovƒõdnost','po≈°kozen√≠','opravy','z√°vady',
    'p≈ôepis','p≈ôepisem','vy√∫ƒçtov√°n√≠','pau≈°√°lnƒõ','roƒçn√≠','mƒõs√≠ƒçn√≠',
    'jena','dominik','ikea','gorenje','bosch','m√∂belix'
}

ROLE_STOP = {
    'pronaj√≠matel','n√°jemce','dlu≈æn√≠k','vƒõ≈ôitel','objednatel','zhotovitel',
    'zamƒõstnanec','zamƒõstnavatel','ruƒçitel','spoludlu≈æn√≠k','jednatel','svƒõdek',
    'statut√°rn√≠','z√°stupce','pojistn√≠k','poji≈°tƒõn√Ω','odes√≠latel','p≈ô√≠jemce',
    'elekt≈ôina','vodn√©','stoƒçn√©','topen√≠','internet','slu≈æba','slu≈æby'
}

# =============== Inference nominativu ===============
def _male_genitive_to_nominative(obs: str) -> Optional[str]:
    lo = obs.lower()
    cands = []
    if lo.endswith('ka') and len(obs) > 2:
        cands.append(obs[:-2] + 'ek')
    if lo.endswith('la') and len(obs) > 2:
        cands.append(obs[:-2] + 'el')
    if lo.endswith('a') and len(obs) > 1:
        cands.append(obs[:-1])
    for cand in cands:
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
    return None

def infer_first_name_nominative(observed: str, surname_observed: str = "") -> Optional[str]:
    if not observed: return None
    obs = observed.strip()
    surname_lower = (surname_observed or "").lower()
    female_like_surname = surname_lower.endswith(('ov√°', '√°', 'ou', '√©'))

    if not female_like_surname:
        cand = _male_genitive_to_nominative(obs)
        if cand: return cand

    norm = normalize_for_matching(obs)
    if norm in CZECH_FIRST_NAMES:
        return obs

    # Speci√°ln√≠ pravidla pro -ice, -≈ôe
    if obs.lower().endswith('ice') and len(obs) > 3:
        cand = obs[:-3] + 'ika'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
        cand = obs[:-3] + 'a'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand
    
    if obs.lower().endswith('≈ôe') and len(obs) > 2:
        cand = obs[:-2] + 'ra'
        if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
            return cand

    for suf in ['inou','in√©','inu','iny','ou','u','y','e','ƒõ','o']:
        if obs.lower().endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)] + 'a'
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand

    for suf in ['ovi','em','e','u']:
        if obs.lower().endswith(suf) and len(obs) > len(suf)+1:
            cand = obs[:-len(suf)]
            if normalize_for_matching(cand) in CZECH_FIRST_NAMES:
                return cand
    return None

def infer_surname_nominative(observed: str) -> str:
    if not observed: return observed
    obs = observed.strip()
    low = obs.lower()

    if low.endswith('ovou') and len(obs) > 4: return obs[:-4] + 'ov√°'
    if low.endswith('ov√©') and len(obs) > 3:  return obs[:-3] + '√°'
    if low.endswith('√©') and len(obs) > 2:    return obs[:-1] + '√°'
    if low.endswith('ou') and not low.endswith('ovou') and len(obs) > 2:
        return obs[:-2] + '√°'

    m = re.match(r'^(.*)ƒçek(a|ovi|em|u|e|y|ou|≈Øm|√°ch)?$', obs, flags=re.IGNORECASE)
    if m: return m.group(1) + 'ƒçek'
    
    m2 = re.match(r'^(.*)nk(a|ovi|em|u|e|y|ou|≈Øm|√°ch)?$', obs, flags=re.IGNORECASE)
    if m2: return m2.group(1) + 'nek'
    
    if low.endswith(('ka','kovi','kem','ku','ke')) and len(obs) > 3:
        return re.sub(r'k(ovi|em|u|e|a)?$', 'ek', obs, flags=re.IGNORECASE)

    m3 = re.match(r'^(.*)c(e|i|em|≈Ø|√≠ch|≈Øm|ech|emi|u|y)?$', obs, flags=re.IGNORECASE)
    if m3: return m3.group(1) + 'ec'

    if low.endswith('ovi') and len(obs) > 4:  return obs[:-3] + 'a'

    for suf in ('em','e','u','y'):
        if low.endswith(suf) and len(obs) > len(suf)+1:
            return obs[:-len(suf)] + 'a'

    return obs

# =============== Varianty pro nahrazov√°n√≠ ===============
def variants_for_first(first: str) -> set:
    f = first.strip()
    if not f: return {''}
    V = {f, f.lower(), f.capitalize()}
    low = f.lower()
    if low.endswith('a'):
        stem = f[:-1]
        V |= {stem+'y', stem+'e', stem+'ƒõ', stem+'u', stem+'ou', stem+'o'}
        V |= {stem+s for s in ['in','ina','iny','in√©','inu','inou','in√Ωm','in√Ωch']}
        if stem.endswith('tr'):
            V |= {stem[:-1]+'≈ô'+s for s in ['in','ina','iny','in√©','inu','inou','in√Ωm','in√Ωch']}
    else:
        V |= {f+'a', f+'ovi', f+'e', f+'em', f+'u', f+'om'}
        V |= {f+'≈Øv'} | {f+'ov'+s for s in ['a','o','y','ƒõ','√Ωm','√Ωch']}
        if low.endswith('ek'): V.add(f[:-2] + 'ka')
        if low.endswith('el'): V.add(f[:-2] + 'la')
    V |= {unicodedata.normalize('NFKD', v).encode('ascii','ignore').decode('ascii') for v in list(V)}
    return V

def variants_for_surname(surname: str) -> set:
    s = surname.strip()
    if not s: return {''}
    out = {s, s.lower(), s.capitalize()}
    low = s.lower()

    if low.endswith('ov√°'):
        base = s[:-1]
        out |= {s, base+'√©', base+'ou'}
        return out
    if low.endswith(('sk√Ω','ck√Ω','√Ω')):
        stem = s[:-1] if low.endswith('√Ω') else s[:-3]
        out |= {stem+'√Ω', stem+'√©ho', stem+'√©mu', stem+'√Ωm', stem+'√©m', stem+'√°', stem+'√©', stem+'ou'}
        return out
    if low.endswith('√°'):
        stem = s[:-1]; out |= {s, stem+'√©', stem+'ou'}; return out
    if low.endswith('ek') and len(s) >= 3:
        stem_k = s[:-2] + 'k'
        out |= {s, stem_k+'a', stem_k+'ovi', stem_k+'em', stem_k+'u', stem_k+'e', stem_k+'y', stem_k+'ou'}
        return out
    if low.endswith('ec') and len(s) >= 3:
        stem_c = s[:-2] + 'c'
        out |= {s, stem_c+'e', stem_c+'i', stem_c+'em', stem_c+'≈Ø', stem_c+'≈Øm', stem_c+'√≠ch', stem_c+'ech', stem_c+'emi', stem_c+'u', stem_c+'y'}
        return out
    if low.endswith('a') and len(s) >= 2:
        stem = s[:-1]
        out |= {s, stem+'y', stem+'ovi', stem+'ou', stem+'u', stem+'e'}
        return out
    out |= {s+'a', s+'ovi', s+'e', s+'em', s+'u'}
    return out

# =============== Regexy ===============
ADDRESS_RE = re.compile(r'(?<!\[)\b[A-Z√Åƒåƒé√âƒö√ç≈á√ì≈ò≈†≈§√ö≈Æ√ù≈Ω][^\n\r,\[\]]{2,50}?\s+\d{1,4}(?:/\d{1,4})?,[ \t]*\d{3}[ \t]?\d{2}[ \t]+[A-Z√Åƒåƒé√âƒö√ç≈á√ì≈ò≈†≈§√ö≈Æ√ù≈Ω][^\n\r,\[\]]{1,40}\b', re.UNICODE)
ACCT_RE    = re.compile(r'\b(?:\d{1,6}-)?\d{2,10}/\d{4}\b')
BIRTHID_RE = re.compile(r'\b\d{6}\s*/\s*\d{3,4}\b')
IDCARD_RE  = re.compile(r'\b\d{6,9}/\d{3,4}\b|\b\d{9}\b|[A-Z]{2,3}[ \t]?\d{6,9}\b')
PHONE_RE   = re.compile(r'(?<!\d)(?:\+420|00420)?[ \t\-]?\d{3}[ \t\-]?\d{3}[ \t\-]?\d{3}(?!\s*/\d{4})\b')
EMAIL_RE   = re.compile(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}')
DATE_RE    = re.compile(r'\b\d{1,2}\.\s*\d{1,2}\.\s*\d{4}\b')
STATUTE_RE = re.compile(r'\b(Sb\.?|z√°kon(a|u)?|z√°kon\s*ƒç\.)\b', re.IGNORECASE)
PAIR_RE    = re.compile(r'(?<!\w)([A-Z√Åƒåƒé√âƒö√ç≈á√ì≈ò≈†≈§√ö≈Æ√ù≈Ω][a-z√°ƒçƒè√©ƒõ√≠≈à√≥≈ô≈°≈•√∫≈Ø√Ω≈æ]{1,})\s+([A-Z√Åƒåƒé√âƒö√ç≈á√ì≈ò≈†≈§√ö≈Æ√ù≈Ω][a-z√°ƒçƒè√©ƒõ√≠≈à√≥≈ô≈°≈•√∫≈Ø√Ω≈æ]{1,})(?!\w)')
TITLES_RE  = re.compile(r'\b(Mgr|Ing|Dr|Ph\.?D|RNDr|MUDr|JUDr|PhDr|PaedDr|ThDr|RCDr|MVDr|DiS|Bc|BcA|MBA|LL\.?M|prof|doc)\.?\s+', re.IGNORECASE)

CTX_OP     = re.compile(r'\b(OP|ƒå√≠slo\s+OP|ƒç√≠slo\s+OP|obƒçansk(√Ω|√©ho|√©mu|√©m|√Ωm)|pr≈Økaz|ƒç\.\s*OP)\b', re.IGNORECASE)
CTX_BIRTH  = re.compile(r'\b(rodn[√©e]\s*ƒç[√≠i]slo|Rƒå|rodn[√©e])\b', re.IGNORECASE)
CTX_BANK   = re.compile(r'\b(√∫ƒçet|√∫ƒçtu|√∫ƒçtem|Bankovn√≠\s+√∫ƒçet|bankovn√≠\s+√∫ƒçet|veden[eya].*u|banka|banky|IBAN|ƒç√≠slo\s+√∫ƒçtu)\b', re.IGNORECASE)
CTX_PERSON = re.compile(
    r'(nar\.|narozen|rodn[√©e]\s*ƒç[√≠i]slo|Rƒå|bytem|trval[√©]\s*bydli≈°t[ƒõi]|'
    r'(e-?mail)|tel\.?|telefon|ƒç\.\s*√∫ƒçtu|IBAN|SPZ|Mgr\.|Ing\.|Bc\.|PhDr\.|JUDr\.)',
    re.IGNORECASE
)
CTX_ROLE   = re.compile(r'\b(pronaj[i√≠]matel|n[a√°]jemce|dlu[z≈æ]n[i√≠]k|v[eƒõ]≈ôitel|objednatel|zhotovitel|zam[eƒõ]stnanec|zam[eƒõ]stnavatel|ruƒçitel|spoludlu[z≈æ]n[i√≠]k|jednatel|statut[a√°]rn[i√≠]\s+z[a√°]stupce|sv[eƒõ]dek)\b', re.IGNORECASE)
CTX_LABEL  = re.compile(r'j[mn][eƒõ]no\s*(,|a)?\s*p≈ô[i√≠]jmen[i√≠]', re.IGNORECASE)

def looks_like_firstname(token: str) -> bool:
    if not token or not token[0].isupper(): return False
    norm = normalize_for_matching(token)
    if norm in CZECH_FIRST_NAMES: return True
    return any([
        norm.endswith('ek'), norm.endswith('el'), norm.endswith('os'),
        norm.endswith('as'), norm.endswith('an'), norm.endswith('en'),
        norm.endswith('a') and len(norm) > 3,
    ])

# =============== Anonymizer ===============
class Anonymizer:
    def __init__(self, verbose=False):
        self.verbose = verbose
        self.counter = defaultdict(int)
        self.tag_map = defaultdict(list)
        self.value_to_tag = {}
        self.person_index = {}
        self.canonical_persons = []
        self.person_variants = {}
        self.source_text = ""

    def _get_or_create_tag(self, cat: str, value: str) -> str:
        norm_val = ' '.join(value.split())
        lookup_key = f"{cat}:{norm_val}"
        if lookup_key in self.value_to_tag:
            return self.value_to_tag[lookup_key]
        self.counter[cat] += 1
        tag = f'[[{cat}_{self.counter[cat]}]]'
        self.value_to_tag[lookup_key] = tag
        self._record_value(tag, value)
        return tag

    def _record_value(self, tag: str, value: str):
        if value and re.search(r'(?<!\w)'+re.escape(value)+r'(?!\w)', self.source_text):
            if value not in self.tag_map[tag]:
                self.tag_map[tag].append(value)

    def _ensure_person_tag(self, first_nom: str, last_nom: str) -> str:
        key = (normalize_for_matching(first_nom), normalize_for_matching(last_nom))
        if key in self.person_index:
            return self.person_index[key]
        tag = self._get_or_create_tag('PERSON', f'{first_nom} {last_nom}')
        self.person_index[key] = tag
        self.canonical_persons.append({'first': first_nom, 'last': last_nom, 'tag': tag})
        fvars = variants_for_first(first_nom)
        svars = variants_for_surname(last_nom)
        self.person_variants[tag] = {f'{f} {s}' for f in fvars for s in svars}
        return tag

    def _extract_persons_to_index(self, text: str):
        text_no_titles = TITLES_RE.sub('', text)
        for m in PAIR_RE.finditer(text_no_titles):
            s, e = m.span()
            f_tok, l_tok = m.group(1), m.group(2)

            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue
            if normalize_for_matching(f_tok) in SURNAME_BLACKLIST:
                continue
            
            pre = text[max(0, s-80):s]
            post = text[e:e+80]
            if re.search(r'\b(v√Ωrobce|model|znaƒçka|invent√°≈ô|v√Ωrobek|polo≈æk)', pre+post, re.IGNORECASE):
                if (normalize_for_matching(f_tok) in SURNAME_BLACKLIST or 
                    normalize_for_matching(l_tok) in SURNAME_BLACKLIST):
                    continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            l_nom = infer_surname_nominative(l_tok)

            if normalize_for_matching(f_nom) in CZECH_FIRST_NAMES:
                self._ensure_person_tag(f_nom, l_nom)
                continue

            pre = text[max(0, s-160):s]
            post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)
            if (has_ctx
                and f_tok[:1].isupper() and l_tok[:1].isupper()
                and looks_like_firstname(f_tok)
                and f_tok.lower() not in ROLE_STOP and l_tok.lower() not in ROLE_STOP):
                self._ensure_person_tag(f_nom, l_nom)

    def _apply_known_people(self, text: str) -> str:
        for p in self.canonical_persons:
            tag = self._ensure_person_tag(p['first'], p['last'])
            for pat in sorted(self.person_variants[tag], key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(pat)+r'(?!\w)', re.IGNORECASE)
                def repl(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl, text)
            
            first_low, last_low = p['first'].lower(), p['last'].lower()
            poss = set()
            if first_low.endswith('a'):
                stem = p['first'][:-1]
                poss |= {stem+s for s in ['in','ina','iny','in√©','inu','inou','in√Ωm','in√Ωch']}
                if stem.endswith('tr'):
                    poss |= {stem[:-1]+'≈ô'+s for s in ['in','ina','iny','in√©','inu','inou','in√Ωm','in√Ωch']}
            else:
                poss |= {p['first']+'≈Øv'} | {p['first']+'ov'+s for s in ['a','o','y','ƒõ','√Ωm','√Ωch']}
            if not last_low.endswith('ov√°'):
                poss |= {p['last']+'≈Øv'} | {p['last']+'ov'+s for s in ['a','o','y','ƒõ','√Ωm','√Ωch']}
            for token in sorted(list(poss), key=len, reverse=True):
                rx = re.compile(r'(?<!\w)'+re.escape(token)+r'(?!\w)', re.IGNORECASE)
                def repl2(m):
                    surf = m.group(0)
                    self._record_value(tag, surf)
                    return preserve_case(surf, tag)
                text = rx.sub(repl2, text)
        return text

    def _replace_remaining_people(self, text: str) -> str:
        text_no_titles = TITLES_RE.sub('', text)
        offset = 0
        for m in list(PAIR_RE.finditer(text_no_titles)):
            s, e = m.start()+offset, m.end()+offset
            seg = text[s:e]
            if seg.startswith('[[') and seg.endswith(']]'):
                continue
            f_tok, l_tok = m.group(1), m.group(2)

            if f_tok.lower() in ROLE_STOP or l_tok.lower() in ROLE_STOP:
                continue
            if normalize_for_matching(l_tok) in SURNAME_BLACKLIST:
                continue

            f_nom = infer_first_name_nominative(f_tok, l_tok) or f_tok
            pre = text[max(0, s-160):s]
            post = text[e:e+160]
            has_ctx = CTX_PERSON.search(pre+post) or CTX_ROLE.search(pre+post) or CTX_LABEL.search(pre+post)

            if (normalize_for_matching(f_nom) not in CZECH_FIRST_NAMES
                and not (has_ctx and looks_like_firstname(f_tok))):
                continue

            l_nom = infer_surname_nominative(l_tok)
            tag = self._ensure_person_tag(f_nom, l_nom)
            before = text
            text = text[:s] + preserve_case(seg, tag) + text[e:]
            self._record_value(tag, seg)
            offset += len(text) - len(before)
        return text

    def _is_statute(self, text: str, s: int, e: int) -> bool:
        pre = text[max(0, s-20):s]
        post = text[e:e+10]
        return bool(STATUTE_RE.search(pre) or STATUTE_RE.search(post))

    def _replace_entity(self, text: str, rx: re.Pattern, cat: str) -> str:
        def repl(m):
            v = m.group(0)
            tag = self._get_or_create_tag(cat, v)
            self._record_value(tag, v)
            return tag
        return rx.sub(repl, text)

    def anonymize_entities(self, text: str) -> str:
        text = self._replace_entity(text, EMAIL_RE, 'EMAIL')

        def addr_repl(m):
            v = m.group(0).strip()
            v = re.sub(r'^(Trval√©\s+bydli≈°tƒõ|Bydli≈°tƒõ|Adresa)\s*:\s*', '', v, flags=re.IGNORECASE)
            v = re.sub(r'^.{0,30}?\b(na\s+adrese|v\s+domƒõ|domu)\s+', '', v, flags=re.IGNORECASE)
            v = re.sub(r'\s*\(d√°le\s+jen.*$', '', v, flags=re.IGNORECASE)
            v = v.strip()
            if not v:
                return m.group(0)
            tag = self._get_or_create_tag('ADDRESS', v)
            self._record_value(tag, v)
            return tag
        text = ADDRESS_RE.sub(addr_repl, text)

        text = self._replace_entity(text, DATE_RE, 'DATE')

        def phone_repl(m):
            v = m.group(0)
            s, e = m.span()
            pre = text[max(0, s-15):s]
            if re.search(r'(OP|obƒçansk\w+|ƒç\.\s*OP)', pre, re.IGNORECASE):
                tag = self._get_or_create_tag('ID_CARD', v)
                self._record_value(tag, v)
                return tag
            if re.match(r'^\s*/\d{4}', text[e:e+6]):
                return v
            tag = self._get_or_create_tag('PHONE', v)
            self._record_value(tag, v)
            return tag
        text = PHONE_RE.sub(phone_repl, text)

        def acct_like(m):
            s, e = m.span()
            if self._is_statute(text, s, e):
                return m.group(0)
            raw = m.group(0)
            
            parts = raw.split('/')
            if len(parts) == 2:
                main_part = parts[0].replace('-', '')
                bank_code = parts[1]
                
                if len(main_part) >= 7 and len(bank_code) == 4:
                    tag = self._get_or_create_tag('BANK', raw)
                    self._record_value(tag, raw)
                    return tag
            
            pre = text[max(0, s-30):s]
            post = text[e:e+30]
            if CTX_BANK.search(pre+post):
                tag = self._get_or_create_tag('BANK', raw)
                self._record_value(tag, raw)
                return tag
            if CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', raw)
                self._record_value(tag, raw)
                return tag
            
            return raw
        text = ACCT_RE.sub(acct_like, text)

        def birth_or_id_repl(m):
            v = m.group(0)
            s, e = m.span()
            pre = text[max(0, s-40):s]
            post = text[e:e+40]
            
            if CTX_OP.search(pre+post):
                tag = self._get_or_create_tag('ID_CARD', v)
            elif CTX_BIRTH.search(pre+post):
                tag = self._get_or_create_tag('BIRTH_ID', v)
            else:
                tag = self._get_or_create_tag('BIRTH_ID', v)
            
            self._record_value(tag, v)
            return tag
        text = BIRTHID_RE.sub(birth_or_id_repl, text)

        def id_repl(m):
            v = m.group(0)
            tag = self._get_or_create_tag('ID_CARD', v)
            self._record_value(tag, v)
            return tag
        text = IDCARD_RE.sub(id_repl, text)

        return text

    def post_merge_person_tags(self, doc: Document):
        key_to_tags = defaultdict(set)
        for tag, vals in list(self.tag_map.items()):
            if not tag.startswith('[[PERSON_'):
                continue
            for v in vals:
                m = PAIR_RE.search(v)
                if not m:
                    continue
                f_nom = infer_first_name_nominative(m.group(1), m.group(2)) or m.group(1)
                l_nom = infer_surname_nominative(m.group(2))
                key = (normalize_for_matching(f_nom), normalize_for_matching(l_nom))
                key_to_tags[key].add(tag)

        redirect = {}
        for key, tags in key_to_tags.items():
            if len(tags) <= 1:
                continue
            canon = sorted(tags)[0]
            for t in tags:
                if t != canon:
                    redirect[t] = canon

        if redirect:
            for p in iter_paragraphs(doc):
                txt = get_text(p)
                new = txt
                for src, dst in redirect.items():
                    new = new.replace(src, dst)
                if new != txt:
                    set_text(p, new)

            for src, dst in redirect.items():
                if src in self.tag_map:
                    for v in self.tag_map[src]:
                        if v not in self.tag_map[dst]:
                            self.tag_map[dst].append(v)
                    del self.tag_map[src]

    def anonymize_docx(self, input_path: str, output_path: str, json_map: str, txt_map: str):
        doc = Document(input_path)
        pieces = []
        for p in iter_paragraphs(doc):
            pieces.append(clean_invisibles(get_text(p)))
        self.source_text = '\n'.join(pieces)

        self._extract_persons_to_index(self.source_text)

        for p in iter_paragraphs(doc):
            raw = get_text(p)
            if not raw.strip():
                continue
            txt = clean_invisibles(raw)
            txt = self._apply_known_people(txt)
            txt = self._replace_remaining_people(txt)
            txt = self.anonymize_entities(txt)
            if txt != raw:
                set_text(p, txt)

        self.post_merge_person_tags(doc)

        doc.save(output_path)

        data = OrderedDict((tag, self.tag_map[tag]) for tag in sorted(self.tag_map.keys()))
        with open(json_map, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        with open(txt_map, 'w', encoding='utf-8') as f:
            sections = [
                ("OSOBY", "PERSON"),
                ("RODN√Å ƒå√çSLA", "BIRTH_ID"),
                ("BANKOVN√ç √öƒåTY", "BANK"),
                ("TELEFONY", "PHONE"),
                ("EMAILY", "EMAIL"),
                ("OBƒåANSK√â PR≈ÆKAZY", "ID_CARD"),
                ("DATA", "DATE"),
                ("ADRESY", "ADDRESS"),
            ]
            for title, pref in sections:
                items = []
                for tag, vals in sorted(self.tag_map.items()):
                    if tag.startswith(f'[[{pref}_'):
                        for v in vals:
                            items.append(f"{tag}: {v}")
                if items:
                    f.write(f"{title}\n{'-'*len(title)}\n")
                    f.write("\n".join(items) + "\n\n")

def main():
    import argparse
    ap = argparse.ArgumentParser(description="Anonymizace ƒçesk√Ωch DOCX s JSON knihovnou jmen")
    ap.add_argument("docx_path", nargs='?', help="Cesta k .docx souboru")
    ap.add_argument("--names-json", default="cz_names.v1.json", help="Cesta k JSON knihovnƒõ jmen")
    args = ap.parse_args()

    if args.names_json != "cz_names.v1.json":
        global CZECH_FIRST_NAMES
        CZECH_FIRST_NAMES = load_names_library(args.names_json)

    path = Path(args.docx_path) if args.docx_path else Path(input("P≈ôet√°hni sem .docx soubor nebo napi≈° cestu: ").strip().strip('"'))
    if not path.exists():
        print("‚ùå Soubor nenalezen:", path)
        return 2
    
    base = path.stem
    out_docx = path.parent / f"{base}_anon.docx"
    out_json = path.parent / f"{base}_map.json"
    out_txt  = path.parent / f"{base}_map.txt"
    
    print(f"\nüîç Zpracov√°v√°m: {path.name}")
    a = Anonymizer(verbose=False)
    a.anonymize_docx(str(path), str(out_docx), str(out_json), str(out_txt))
    
    print("\n‚úÖ V√Ωstupy:")
    print(f" - {out_docx}")
    print(f" - {out_json}")
    print(f" - {out_txt}")
    print(f"\nüìä Statistiky:")
    print(f" - Nalezeno osob: {len(a.canonical_persons)}")
    print(f" - Celkem tag≈Ø: {sum(a.counter.values())}")

if __name__ == "__main__":
    sys.exit(main())